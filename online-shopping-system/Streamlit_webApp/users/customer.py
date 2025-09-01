from users.user import User
from models.product import Product
from models.cart import ShoppingCart
from models.payment import Payment
from datetime import datetime
import csv
import streamlit as st
import pandas as pd
import io
from docx import Document
     
# customer class inherit from user class
class Customer(User):
    '''Customer class for shopping'''
    def __init__(self,name):
        '''Initialize customer with name and role'''
        super().__init__(name,role='Customer')
        self.cart=ShoppingCart()
    def add_to_cart(self,product_name,Quantity):
        '''Add product to cart'''
        try:
           with open('Streamlit-App/data/products.csv','r') as f:
                  reader=csv.DictReader(f)
                  products=[row for row in reader if row]
        except FileNotFoundError:
           st.error('File Not Found')
           return
        products_found=False
        for row in products:
            if row['product_name'].strip().lower()==product_name.strip().lower():
                products_found=True
                price=float(row['price'])
                stock=int(row['stock'])
                if stock >= Quantity :
                    self.cart.add_product(product_name,price,Quantity)
                    st.success(f'{Quantity} x {product_name} added to cart at ${price} each!')
                else:
                    st.error(f'Not enough stock for "{product_name}", Available stock: {stock}.')
                return
        if not products_found:
            st.error(f'Product "{product_name}" not found .')
    def remove_to_cart(self,product_name):
        '''Remove product from cart'''
        self.cart - product_name
    def view_cart(self):
        '''View cart'''
        if not self.cart.items:
            st.error('Your cart is empty.')
            return
        Total_price=self.cart.get_total_price()
        time=datetime.now().strftime("**Date**: %y-%m-%d & **Time**: %H:%M:%S")
        st.subheader(f"{self.name.upper()}'s CART 🛒")
        st.write('='*30)
        st.write(f"{time}")
        st.write(f"**Customer Name** : {self.name}")
        st.write(f"**Total Price** : ${Total_price}")
        st.write('='*30)
        st.write('Items Details :')

        df=[]
        for product,details in self.cart.items.items():
            df.append([product,details['Quantity'],details['Price']])
        df=pd.DataFrame(df,columns=['Product Name','Quantity','Price'])
        st.table(df)
    
    def checkout(self):
        '''Checkout'''
        if not self.cart.items:
            st.error('Your cart is empty.')
            return
        total_price=self.cart.get_total_price()
        payment=Payment(total_price,self.name)
        time=datetime.now().strftime("**Date** : %y-%m-%d & **Time** : %H:%M:%S")
        if payment.process_payment():
            st.session_state['checkout_success']=True
            st.session_state['cart_items_copy']=self.cart.items.copy()
            st.session_state['payment_details']='.'.join(payment.details)
            st.session_state['payment_method']=payment.payment_method
            st.session_state['total_price']=total_price
            st.session_state['time']=time
            #update stock
            for product_name,details in self.cart.items.items():
               quantity=int(details['Quantity'])
               Product.update_stock(product_name,quantity)
            self.cart.clear_cart()
        else:
            return
    def generate_receipt(self):
        '''Generate receipt'''
        if not st.session_state.get('checkout_success'):
            return
        cart_items=st.session_state.get('cart_items_copy',{})
        time=st.session_state.get('time','')
        total_price=st.session_state.get('total_price',0)
        payment_method=st.session_state.get('payment_method')
        payment_details=st.session_state.get('payment_details')
        
        st.subheader("RECEIPT 📝")
        st.write('='*30)
        st.write(f"{time}")
        st.write(f"**Customer Name** : {self.name}")
        st.write(f"**Total Price** : ${total_price}")
        st.write(f"**Payment Method** : {payment_method}")
        st.write(f"**Payment Details** : {payment_details}")
        st.write('='*30)
        st.write('Purchased Items Details :')
        df=[]
        for product,details in cart_items.items():
            df.append([product,details['Quantity'],details['Price']])
        df=pd.DataFrame(df,columns=['Product Name','Quantity','Price'])
        st.table(df)
        st.write("Thank you for shopping with us! 🛒 ")

        #download receipt
        file_name,file_buffer =self.download_receipt(time,total_price,payment_method,payment_details,cart_items)
        st.download_button(
                label='Download Receipt',
                key='download_receipt',
                data=file_buffer,
                file_name=file_name,
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        st.info(f'Receipt Generated successfully as "{file_name}".')

    def download_receipt(self,time,total_price,payment_method,payment_details,cart_items):
        '''Download receipt'''          
        receipt_time=datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        file_name=f"{self.name}_receipt_{receipt_time}.docx"
        buffer=io.BytesIO()
        doc=Document()
        doc.add_heading('Receipt',0)
        doc.add_paragraph(f"**Customer Name** : {self.name}")
        doc.add_paragraph(f"**Total Price** : ${total_price}")
        doc.add_paragraph(f"Time : {time}")
        doc.add_paragraph('-'*40)
        doc.add_paragraph('Purchased Items Details :')
        for product,details in cart_items.items():
            doc.add_paragraph(f"{product} x {details['Quantity']} - ${details['Price']} each")
        doc.add_paragraph('-'*40)
        doc.add_paragraph(f"**Payment Method** : {payment_method}")
        doc.add_paragraph(f"**Payment Details** : {payment_details}")
        doc.add_paragraph(f"Thank you for shopping with us! 🛒 ")
        doc.save(buffer)
        buffer.seek(0)
        return file_name,buffer
        
